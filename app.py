import os
import sqlite3
from io import BytesIO

from docx import Document


from flask import Flask, render_template, redirect, request, flash, send_from_directory,send_file
from werkzeug.exceptions import abort

app = Flask(__name__)
app.config['SECRET_KEY'] = b'my)secret)key'
UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def get_db_connection():
    conn = sqlite3.connect('../pythonProject4/database.db')
    conn.row_factory = sqlite3.Row
    return conn


@app.route('/')
def index():
    return redirect("/contracts")


@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    """ Страница-добавления нового контракта """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы

        number = request.form['number']
        date = request.form['date']
        price = int(request.form['price'])
        discount = int(request.form['discount'])
        finish_price = price * (1 - discount * 0.01)
        deal_status = 0
        client_id = int(request.form.get('client'))
        training_id = int(request.form.get('training'))
        employee_id = int(request.form.get('employee'))
        if not (client_id > 0 and training_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (number and date and price and discount):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'contracts' ('number', 'date', 'price', 'discount', 'deal_status', 'finish_price', 'training_id', 'client_id', 'employee_id')  VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    (number, date, price, discount, deal_status, finish_price, training_id, client_id,
                     employee_id))
                conn.commit()
                new_contract_id = cursor.lastrowid
                print(new_contract_id)
                conn.close()
                return redirect(f'/contract/{new_contract_id}')

    # отрисовка формы
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos2 = conn.execute("""SELECT * FROM trainings""").fetchall()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('new_contract.html', clients=pos1, trainings=pos2, employees=pos3)


@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    """ Страница генерации договора """

    # переменные шаблона
    id = int(request.args.get('id_contract'))
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, trainings, clients, employees
            WHERE contracts.training_id = trainings.id_training and 
                        contracts.client_id = clients.id_client and 
                        contracts.employee_id = employees.id_employee and
                        contracts.id_contract = ?
                        """, (id,)).fetchone()
    conn.close()
    contract_params = {
        'CLIENT_BIRTHDATE': 'дата рождения клиента',
        'CLIENT_BIRTHPLACE': 'место рождения клиента',
        'CLIENT_PASSPORT_DATE': 'дата выдачи паспорта клиента',
        'CLIENT_PASSPORT_DEPARTMENT': 'подразделение, выдавшее паспорт клиента',
        'CLIENT_PASSPORT_DEPCODE': 'код подразделения, выдавшего паспорт клиента',
        'CLIENT_REG_ADDRESS': 'адрес регистрации клиента',
        'CONTRACT_DURATION': 'длительность контракта',
        'CONTRACT_PRICE': 'цена абонемента'}
    contract_params_auto = {
        'CONTRACT_NUMBER': ['номер договора', pos['number']],
        'CONTRACT_DATE': ['дата подписания договора', pos['date']],
        'CLIENT_FULLNAME': ['ФИО клиента', pos['name']],
        'CLIENT_PASSPORT_NUMBER': ['серия и номер паспорта клиента', pos['passport']],
        'EMPLOYEE_POSITION': ['должность сотрудника риэлторской компании', pos['position']],
        'EMPLOYEE_FULLNAME': ['ФИО сотрудника риэлторской компании', pos['name']], }

    if request.method == 'POST':
        # создание нового документа
        result_params = request.form.to_dict()
        create_contract(id, result_params)
        filename = f"договор {pos['number']} от {pos['date']}.docx"
        return send_file(filename)

    else:
        # отрисовка формы заполнения
        flash('Договор не сформирован, заполните его')
        return render_template('generate_contract.html',
                               contract=pos, contract_params=contract_params, auto_params=contract_params_auto)


def create_contract(id, contract_params):
    """ Создание нового документа по шаблону """
    result = os.path.join(
                          f"договор {contract_params['CONTRACT_NUMBER']} от {contract_params['CONTRACT_DATE']}.docx")
    template_path = 'contract_template.docx'
    template_doc = Document(template_path)
    for key, value in contract_params.items():
        for paragraph in template_doc.paragraphs:
            replace_text(paragraph, f'=={key}==', value)
        for table in template_doc.tables:
            replace_text_in_tables(table, key, value)
    template_doc.save(result)



def replace_text(paragraph, key, value):
    """ Работа docx - заполнение параграфов """

    if key in paragraph.text:
        paragraph.text = paragraph.text.replace(key, value)


def replace_text_in_tables(table, key, value):
    """ Работа docx - заполнение таблиц """

    for row in table.rows:
        for cell in row.cells:
            if key in cell.text:
                cell.text = cell.text.replace(key, value)


# Контракты

@app.route('/contracts')
def contracts():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM contracts, trainings, clients
    WHERE contracts.training_id = trainings.id_training and contracts.client_id = clients.id_client
    """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=pos)


def get_contract(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM contracts, trainings, clients
    WHERE contracts.training_id = trainings.id_training and contracts.client_id = clients.id_client and contracts.id_contract = ?
                        """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    pos = get_contract(contract_id)
    return render_template('contract.html', contract=pos)


# Квартиры
@app.route('/new_training', methods=('GET', 'POST'))
def new_training():
    """ Страница-добавления новой квартиры """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            duration = int(request.form['duration'])
            description = request.form['description']
            amount = int(request.form['amount'])
            client_id = int(request.form.get('client'))
            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            client_id = 0
        if not (client_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            if not (duration and description and amount):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'trainings' ('duration', 'description', 'amount', 'client_id', 'employee_id')  VALUES (?, ?, ?, ?, ?)",
                    (duration, description, amount, client_id,
                     employee_id))
                conn.commit()
                new_training_id = cursor.lastrowid
                print(new_training_id)
                conn.close()
                return redirect(f'/training/{new_training_id}')
    # отрисовка формы
    conn = get_db_connection()
    pos1 = conn.execute("""SELECT * FROM clients""").fetchall()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('new_training.html', clients=pos1, employees=pos3)


@app.route('/trainings')
def trainings():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM trainings""").fetchall()
    conn.close()
    return render_template('trainings.html', trainings=pos)


def get_training(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM trainings WHERE id_training = ?""", (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/training/<int:training_id>')
def training(training_id):
    pos = get_training(training_id)
    return render_template('training.html', training=pos)


@app.route('/new_client', methods=('GET', 'POST'))
def new_client():
    """ Страница-добавления новой квартиры """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            name = request.form['name']
            email = request.form['email']
            phone_number = request.form['phone_number']
            passport = request.form['passport']
        except ValueError:
            flash('Некорректные значения')
            passport = 0
        else:
            if not (name and email and phone_number):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'clients' ('name', 'email', 'phone_number', 'passport')  VALUES (?, ?, ?, ?)",
                    (name, email, phone_number, passport))
                conn.commit()
                new_client_id = cursor.lastrowid
                print(new_client_id)
                conn.close()
                return redirect(f'/client/{new_client_id}')
    # отрисовка формы
    return render_template('new_client.html')


@app.route('/clients')
def clients():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM clients""").fetchall()
    conn.close()
    return render_template('clients.html', clients=pos)


def get_client(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM clients WHERE id_client = ?""", (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/client/<int:client_id>')
def client(client_id):
    pos = get_client(client_id)
    return render_template('client.html', client=pos)


@app.route('/new_report', methods=('GET', 'POST'))
def new_report():
    """ Страница-добавления новой квартиры """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = request.form['description']
            employee_id = int(request.form.get('employee'))
        except ValueError:
            flash('Некорректные значения')
            return render_template('new_report.html')  # прерываем выполнение функции
        else:
            if not (number and date and report_type and description and employee_id):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'reports' ('number', 'date', 'report_type', 'description', 'employee_id')  VALUES (?, ?, ?, ?, ?)",
                    (number, date, report_type, description,
                     employee_id))
                conn.commit()
                new_report_id = cursor.lastrowid
                print(new_report_id)
                conn.close()
                return redirect(f'/report/{new_report_id}')
    # отрисовка формы
    conn = get_db_connection()
    pos3 = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()  # закрываем соединение
    return render_template('new_report.html', employees=pos3)


@app.route('/reports')
def reports():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM reports""").fetchall()
    conn.close()
    return render_template('reports.html', reports=pos)


def get_report(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM reports WHERE id_report = ?""", (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/report/<int:report_id>')
def report(report_id):
    pos = get_report(report_id)
    return render_template('report.html', report=pos)


@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():
    """ Страница-добавления новой квартиры """

    if request.method == 'POST':
        # добавление нового контракта в БД псоле заполнения формы
        try:
            name = request.form['name']
            email = request.form['email']
            phone_number = request.form['phone_number']
            position = request.form['position']
            department = request.form['department']
            chief_id = int(request.form['chief_id'])
        except ValueError:
            flash('Некорректные значения')
            position = 0
        else:
            if not (name and email and phone_number):
                flash('Не все поля заполнены')
            else:
                conn = get_db_connection()
                cursor = conn.cursor()
                cursor.execute(
                    "INSERT INTO 'employees' ('name', 'email', 'phone_number', 'position','department','chief_id')  VALUES (?, ?, ?, ?, ?, ?)",
                    (name, email, phone_number, position, department, chief_id))
                conn.commit()
                new_employee_id = cursor.lastrowid
                print(new_employee_id)
                conn.close()
                return redirect(f'/employee/{new_employee_id}')
    return render_template('new_employee.html')


@app.route('/employees')
def employees():
    conn = get_db_connection()
    pos = conn.execute("""SELECT * FROM employees""").fetchall()
    conn.close()
    return render_template('employees.html', employees=pos)


def get_employee(item_id):
    conn = get_db_connection()
    item = conn.execute("""SELECT * FROM employees WHERE id_employee = ?""", (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)
    return item


@app.route('/employee/<int:employee_id>')
def employee(employee_id):
    pos = get_employee(employee_id)
    return render_template('employee.html', employee=pos)


@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404


if __name__ == '__main__':
    app.run()
